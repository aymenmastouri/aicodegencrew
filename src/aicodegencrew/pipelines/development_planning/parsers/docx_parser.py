"""
DOCX Parser for Word documents.
"""

from pathlib import Path
from typing import Dict, Any, List


def parse_docx(file_path: Path) -> Dict[str, Any]:
    """
    Parse DOCX file and extract structured content.

    Args:
        file_path: Path to DOCX file

    Returns:
        Dictionary with title, sections, tables
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. "
            "Install with: pip install python-docx"
        )

    doc = Document(file_path)

    result = {
        'title': '',
        'sections': [],
        'tables': [],
    }

    # Extract title (first heading or first paragraph)
    for para in doc.paragraphs:
        if para.text.strip():
            result['title'] = para.text.strip()
            break

    # Extract sections
    current_section = {'title': '', 'content': []}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if heading
        if para.style.name.startswith('Heading'):
            if current_section['content']:
                result['sections'].append(current_section)
            current_section = {'title': text, 'content': []}
        else:
            current_section['content'].append(text)

    if current_section['content']:
        result['sections'].append(current_section)

    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        result['tables'].append(table_data)

    return result
